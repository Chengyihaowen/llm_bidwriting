import os
import re
from datetime import datetime
from flask import Blueprint, request, send_file, current_app
from extensions import db
from models import Project, OutlineNode, ChapterContent, ExportTask, TaskStatus
from utils.response import success, error

bp = Blueprint('export', __name__)


def _apply_run_font(run, font_name='宋体', size=None, bold=False, italic=False, color='000000'):
    """统一设置 run 的字体、颜色等"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size is not None:
        run.font.size = size
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _write_heading_paragraph(doc, text, level):
    """写标题段落，强制黑体黑色，不继承 Word 默认蓝色 Heading 样式"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    if level == 1:
        font_size = Pt(16)
    elif level == 2:
        font_size = Pt(14)
    else:
        font_size = Pt(13)
    run = p.add_run(text)
    run.bold = True
    _apply_run_font(run, font_name='黑体', size=font_size, bold=True, color='000000')
    return p


def _write_normal_paragraph(doc, text):
    """写正文段落，统一宋体小四黑色"""
    from docx.shared import Pt
    p = doc.add_paragraph()
    _add_runs_with_formatting(p, text)
    return p


def _add_runs_with_formatting(paragraph, text: str):
    """处理 **粗体**、*斜体*、`代码` inline 格式，清理残留 markdown 符号"""
    from docx.shared import Pt
    # 先去掉行内代码反引号（转为普通文本）
    text = re.sub(r'`([^`]*)`', r'\1', text)
    # 分割粗体/斜体
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            _apply_run_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            _apply_run_font(run, italic=True)
        else:
            if part:
                run = paragraph.add_run(part)
                _apply_run_font(run)


def _is_table_line(line: str) -> bool:
    """判断是否是 Markdown 表格行（含 | 字符）"""
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and len(s) > 2


def _is_separator_line(line: str) -> bool:
    """判断是否是表格分隔行（如 |---|---|）"""
    s = line.strip()
    if not s.startswith('|') or not s.endswith('|'):
        return False
    inner = s[1:-1]
    cells = inner.split('|')
    return all(re.match(r'^[\s:\-]+$', c) for c in cells)


def _parse_table_row(line: str) -> list:
    """解析一行表格，返回单元格文本列表"""
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [cell.strip() for cell in s.split('|')]


def _add_docx_table(doc, rows: list):
    """将二维字符串列表渲染为 Word 表格，首行为表头"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            is_header = (r_idx == 0)
            _apply_run_font(run, font_name='宋体', size=Pt(11), bold=is_header, color='000000')


def _markdown_to_docx_paragraphs(doc, markdown_text: str, styles: dict, skip_leading_heading: bool = False):
    """将 markdown 文本转换为 docx 段落，支持表格、标题、列表、粗体斜体"""
    from docx.shared import Pt

    lines = markdown_text.split('\n')
    i = 0
    first_heading_skipped = False  # 是否已跳过首个重复标题

    while i < len(lines):
        line = lines[i]

        # 跳过纯分隔线
        if re.match(r'^---+\s*$', line.strip()):
            i += 1
            continue

        # 标题
        is_heading = False
        heading_level = 0
        heading_text = ''
        if line.startswith('### '):
            is_heading, heading_level, heading_text = True, 3, line[4:].strip()
        elif line.startswith('## '):
            is_heading, heading_level, heading_text = True, 2, line[3:].strip()
        elif line.startswith('# '):
            is_heading, heading_level, heading_text = True, 1, line[2:].strip()

        if is_heading:
            # 跳过正文开头与章节树标题重复的第一个标题
            if skip_leading_heading and not first_heading_skipped:
                first_heading_skipped = True
                i += 1
                continue
            _write_heading_paragraph(doc, heading_text, level=heading_level)
            i += 1
            continue

        # Markdown 表格：收集连续的表格行
        if _is_table_line(line):
            table_lines = []
            while i < len(lines) and _is_table_line(lines[i]):
                if not _is_separator_line(lines[i]):
                    table_lines.append(_parse_table_row(lines[i]))
                i += 1
            if table_lines:
                _add_docx_table(doc, table_lines)
            continue

        # 无序列表
        if line.strip().startswith(('- ', '* ', '+ ')):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_runs_with_formatting(p, text)
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            _add_runs_with_formatting(p, text)
            i += 1
            continue

        # 空行
        if line.strip() == '':
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_runs_with_formatting(p, line)
        i += 1


def _build_docx(project: Project, nodes: list, contents_map: dict, cover_fields: dict) -> str:
    """生成Word文档并返回文件路径"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import copy

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # 设置默认字体
    from docx.oxml import OxmlElement
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ========= 封面 =========
    doc.add_paragraph('')
    doc.add_paragraph('')

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(cover_fields.get('projectName', project.name))
    title_run.font.name = '黑体'
    title_run.font.size = Pt(22)
    title_run.bold = True

    doc.add_paragraph('')

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run('投  标  文  件')
    subtitle_run.font.name = '黑体'
    subtitle_run.font.size = Pt(18)

    doc.add_paragraph('')
    doc.add_paragraph('')

    info_items = [
        ('招标编号', cover_fields.get('tenderNo', project.tender_no or '')),
        ('投标单位', cover_fields.get('bidderName', project.bidder_name or '')),
        ('编制日期', cover_fields.get('date', datetime.now().strftime('%Y年%m月%d日'))),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.name = '宋体'
        run.font.size = Pt(14)

    # 分页
    doc.add_page_break()

    # ========= 正文 =========
    def write_node(node, level=1):
        content = contents_map.get(node.id)
        text = content.content if content and content.content else ''

        _write_heading_paragraph(doc, node.title, level=min(level, 3))

        if text:
            _markdown_to_docx_paragraphs(doc, text, {}, skip_leading_heading=True)
        elif not node.children:
            p = doc.add_paragraph('（本节内容待填写）')
            p.runs[0].italic = True

        children = sorted(node.children, key=lambda x: x.order_no)
        for child in children:
            write_node(child, level=level + 1)

    for node in nodes:
        write_node(node)
        if node != nodes[-1]:
            doc.add_page_break()

    # 保存
    export_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], 'exports')
    os.makedirs(export_dir, exist_ok=True)
    filename = f"bid_{project.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    file_path = os.path.join(export_dir, filename)
    doc.save(file_path)
    return file_path


@bp.route('/<int:project_id>/exports', methods=['POST'])
def create_export(project_id):
    project = Project.query.get_or_404(project_id)
    data = request.get_json() or {}
    cover_fields = data.get('coverFields', {})
    template_name = data.get('templateName', 'standard-bid-v1')

    # 获取目录树
    roots = OutlineNode.query.filter_by(
        project_id=project_id, parent_id=None
    ).order_by(OutlineNode.order_no).all()

    if not roots:
        return error('目录为空，无法导出', 400)

    # 获取所有章节内容
    all_contents = ChapterContent.query.filter_by(project_id=project_id).all()
    contents_map = {c.outline_node_id: c for c in all_contents}

    # 创建导出任务
    task = ExportTask(
        project_id=project_id,
        template_name=template_name,
        status=TaskStatus.RUNNING,
    )
    db.session.add(task)
    db.session.commit()

    try:
        file_path = _build_docx(project, roots, contents_map, cover_fields)
        task.status = TaskStatus.SUCCESS
        task.file_path = file_path
        task.finished_at = datetime.utcnow()
        db.session.commit()
        return success(task.to_dict()), 201
    except Exception as e:
        task.status = TaskStatus.FAILED
        task.error_message = str(e)
        task.finished_at = datetime.utcnow()
        db.session.commit()
        return error(f'导出失败: {str(e)}', 500, 4006)


@bp.route('/<int:project_id>/exports', methods=['GET'])
def list_exports(project_id):
    Project.query.get_or_404(project_id)
    tasks = ExportTask.query.filter_by(project_id=project_id).order_by(ExportTask.created_at.desc()).all()
    return success([t.to_dict() for t in tasks])


@bp.route('/<int:project_id>/exports/<int:export_id>/download', methods=['GET'])
def download_export(project_id, export_id):
    task = ExportTask.query.filter_by(project_id=project_id, id=export_id).first_or_404()
    if not task.file_path or not os.path.exists(task.file_path):
        return error('文件不存在或已删除', 404)

    project = Project.query.get(project_id)
    filename = f"{project.name}_投标文件.docx"
    return send_file(
        task.file_path,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )

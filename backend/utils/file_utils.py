import os
from flask import current_app

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_extension(filename):
    return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

def extract_text_from_file(file_path: str, file_type: str) -> str:
    """从文件中提取纯文本"""
    text = ''
    try:
        if file_type == 'pdf':
            text = _extract_pdf(file_path)
        elif file_type in ('docx', 'doc'):
            text = _extract_docx(file_path)
    except Exception as e:
        raise RuntimeError(f'文件文本提取失败: {e}')
    return text

def _extract_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return '\n'.join(pages)
    except Exception:
        import PyPDF2
        text = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return '\n'.join(text)

def _extract_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(' | '.join(cells))
    return '\n'.join(paragraphs)
